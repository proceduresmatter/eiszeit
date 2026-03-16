"""
DSGVO-Dokumentation Eiszeit – Marta und Florian Schmidl
Enthält:
  Seite 1–2  Verzeichnis von Verarbeitungstätigkeiten (VVT) – Art. 30 DSGVO
  Seite 3    Technische und organisatorische Maßnahmen (TOMs)
  Seite 4    Löschkonzept
"""
import sys
sys.path.insert(0, r'C:\Users\localUser\AppData\Roaming\Python\Python314\site-packages')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date

HEUTE = date.today().strftime("%d.%m.%Y")

BEIGE      = RGBColor(0xF0, 0xE4, 0xCB)
BLAU       = RGBColor(0x9B, 0xD4, 0xDB)
GRUEN      = RGBColor(0x6C, 0x94, 0x8C)
SCHRIFT    = RGBColor(0x41, 0x45, 0x45)
WEISS      = RGBColor(0xFF, 0xFF, 0xFF)
DUNKELBLAU = RGBColor(0x1A, 0x5F, 0x7A)
HELLGRAU   = RGBColor(0xF5, 0xF5, 0xF5)

LOGO   = "EiszeitLogo.png"
OUTPUT = "DSGVO_Eiszeit_VVT_TOMs_Loeschkonzept.docx"

# ── Hilfsfunktionen ───────────────────────────────────────────────────────────

def cell_bg(cell, rgb):
    tc = cell._tc
    pr = tc.get_or_add_tcPr()
    s  = OxmlElement("w:shd")
    s.set(qn("w:val"),   "clear")
    s.set(qn("w:color"), "auto")
    s.set(qn("w:fill"),  f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    pr.append(s)

def cell_border(cell, color="9BD4DB", sz=8):
    tc = cell._tc
    pr = tc.get_or_add_tcPr()
    b  = OxmlElement("w:tcBorders")
    for side in ("top","bottom","left","right"):
        e = OxmlElement(f"w:{side}")
        e.set(qn("w:val"),   "single")
        e.set(qn("w:sz"),    str(sz))
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), color)
        b.append(e)
    pr.append(b)

def run(para, text, bold=False, italic=False, size=10, color=None):
    r = para.add_run(text)
    r.bold      = bold
    r.italic    = italic
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return r

def gap(doc, pt=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(pt)

def page_title(doc, nummer, titel, untertitel=""):
    """Farbiger Seitentitel."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.columns[0].width = Cm(17.4)
    c = t.rows[0].cells[0]
    cell_bg(c, DUNKELBLAU)
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run(p, f"{nummer}  {titel}", bold=True, size=16, color=WEISS)
    if untertitel:
        p2 = c.add_paragraph()
        p2.paragraph_format.left_indent  = Cm(0.5)
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after  = Pt(6)
        run(p2, untertitel, size=9, italic=True, color=BLAU)

def section_bar(doc, text, bg=GRUEN):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.columns[0].width = Cm(17.4)
    c = t.rows[0].cells[0]
    cell_bg(c, bg)
    p = c.paragraphs[0]
    p.paragraph_format.left_indent  = Cm(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run(p, text, bold=True, size=10, color=WEISS)
    gap(doc, 1)

def two_col_row(doc, label, value, label_w=Cm(5.0), val_w=Cm(12.4)):
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.columns[0].width = label_w
    t.columns[1].width = val_w
    lc, rc = t.rows[0].cells
    cell_bg(lc, BEIGE);  cell_border(lc)
    cell_bg(rc, WEISS);  cell_border(rc)
    pl = lc.paragraphs[0]
    pl.paragraph_format.left_indent  = Cm(0.3)
    pl.paragraph_format.space_before = Pt(4)
    pl.paragraph_format.space_after  = Pt(4)
    run(pl, label, bold=True, size=9.5, color=GRUEN)
    pr = rc.paragraphs[0]
    pr.paragraph_format.left_indent  = Cm(0.3)
    pr.paragraph_format.space_before = Pt(4)
    pr.paragraph_format.space_after  = Pt(4)
    run(pr, value, size=9.5, color=SCHRIFT)
    gap(doc, 1)

def check_table(doc, rows_data, col_widths):
    """Tabelle mit Checkboxen (Ja/Nein/Beschreibung)."""
    t = doc.add_table(rows=1, cols=len(col_widths))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, w in enumerate(col_widths):
        t.columns[i].width = w
    # Kopfzeile
    headers = ["Maßnahme", "Umgesetzt", "Beschreibung"]
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        cell_bg(c, GRUEN)
        cell_border(c, sz=6)
        p = c.paragraphs[0]
        p.paragraph_format.left_indent  = Cm(0.2)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        run(p, h, bold=True, size=9, color=WEISS)
    # Datenzeilen
    for row_data in rows_data:
        row = t.add_row()
        bgs = [WEISS, HELLGRAU, WEISS]
        for i, (val, bg) in enumerate(zip(row_data, bgs)):
            c = row.cells[i]
            cell_bg(c, bg)
            cell_border(c, sz=4)
            p = c.paragraphs[0]
            p.paragraph_format.left_indent  = Cm(0.2)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            bold = (i == 1)
            color = GRUEN if (i == 1 and val == "Ja") else SCHRIFT
            run(p, val, bold=bold, size=9, color=color)
    gap(doc, 4)

def new_page(doc):
    doc.add_section(WD_SECTION.NEW_PAGE)

def logo_header(doc, doc_title, art_ref):
    """Kleine Kopfzeile mit Logo + Dokumenttitel."""
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.columns[0].width = Cm(3.0)
    t.columns[1].width = Cm(14.4)
    lc, tc = t.rows[0].cells
    cell_bg(lc, BEIGE); cell_bg(tc, BEIGE)
    lp = lc.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lp.paragraph_format.space_before = Pt(4)
    try:
        lp.add_run().add_picture(LOGO, width=Cm(2.4))
    except Exception:
        run(lp, "EISZEIT", bold=True, size=12, color=DUNKELBLAU)
    tp = tc.paragraphs[0]
    tp.paragraph_format.left_indent  = Cm(0.4)
    tp.paragraph_format.space_before = Pt(6)
    run(tp, "EISZEIT  ·  Marta und Florian Schmidl  ·  Heiligenbruch 27, 27339 Riede",
        bold=True, size=9, color=DUNKELBLAU)
    tp2 = tc.add_paragraph()
    tp2.paragraph_format.left_indent = Cm(0.4)
    run(tp2, f"{doc_title}  ·  {art_ref}  ·  Stand: {HEUTE}",
        size=8, italic=True, color=SCHRIFT)
    gap(doc, 6)

def footer_line(doc):
    gap(doc, 4)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, f"Eiszeit · Marta und Florian Schmidl · Heiligenbruch 27, 27339 Riede · eiszeit@mailbox.org · Stand: {HEUTE}",
        size=7.5, italic=True, color=RGBColor(0xAA,0xAA,0xAA))


# ═══════════════════════════════════════════════════════════════════════════════
#  DOKUMENT
# ═══════════════════════════════════════════════════════════════════════════════
doc = Document()

def set_a4(sec):
    sec.page_width    = Cm(21.0)
    sec.page_height   = Cm(29.7)
    sec.top_margin    = Cm(1.5)
    sec.bottom_margin = Cm(1.5)
    sec.left_margin   = Cm(1.8)
    sec.right_margin  = Cm(1.8)

set_a4(doc.sections[0])


# ═══════════════════════════════════════════════════════════════════════════════
#  SEITE 1+2: VVT
# ═══════════════════════════════════════════════════════════════════════════════
logo_header(doc, "Verzeichnis von Verarbeitungstätigkeiten (VVT)", "Art. 30 DSGVO")
page_title(doc, "A", "Verzeichnis von Verarbeitungstätigkeiten",
           "Art. 30 Abs. 1 DSGVO  ·  Verantwortlicher: Marta und Florian Schmidl")
gap(doc, 6)

# ── VVT-Eintrag 1: Videoüberwachung ──────────────────────────────────────────
section_bar(doc, "Eintrag 1 – Videoüberwachung des Eisautomaten-Standorts")

two_col_row(doc, "Bezeichnung:",
    "Videoüberwachung des Standorts Heiligenbruch 27, 27339 Riede")
two_col_row(doc, "Verantwortlicher:",
    "Marta und Florian Schmidl\nHeiligensbruch 27, 27339 Riede\neizeit@mailbox.org")
two_col_row(doc, "Datenschutzbeauftragter:",
    "Nicht bestellt (Voraussetzungen gem. Art. 37 DSGVO nicht erfüllt)")
two_col_row(doc, "Zweck der Verarbeitung:",
    "Vandalismusprävention, Schutz des Eigentums und der Vertrauenskasse, Hausrecht")
two_col_row(doc, "Rechtsgrundlage:",
    "Art. 6 Abs. 1 lit. f DSGVO (berechtigtes Interesse)\n"
    "Berechtigtes Interesse: Schutz des Eigentums und der Vertrauenskasse "
    "vor Vandalismus und Diebstahl")
two_col_row(doc, "Kategorien betroffener Personen:",
    "Besucher und Kunden des Eisautomaten-Standorts")
two_col_row(doc, "Kategorien personenbezogener Daten:",
    "Bilddaten (Videoaufnahmen), ggf. Bewegungsprofile im überwachten Bereich")
two_col_row(doc, "Empfänger / Kategorien von Empfängern:",
    "Grundsätzlich keine Weitergabe an Dritte.\n"
    "Im begründeten Einzelfall: Strafverfolgungsbehörden (Polizei, Staatsanwaltschaft)")
two_col_row(doc, "Übermittlung in Drittländer:",
    "Keine Übermittlung in Drittländer oder an internationale Organisationen")
two_col_row(doc, "Speicherdauer:",
    "7 Tage; automatische Überschreibung/Löschung nach Ablauf der Frist\n"
    "Ausnahme: Bei dokumentiertem Vorfall ggf. längere Aufbewahrung bis zur Klärung")
two_col_row(doc, "Technische und organisatorische Maßnahmen:",
    "Siehe Anlage TOMs (Seite 3 dieses Dokuments)")
two_col_row(doc, "Zuletzt aktualisiert:", HEUTE)

gap(doc, 8)
section_bar(doc, "Eintrag 2 – Lieferanten- und Belegnachweise (Rückverfolgbarkeit)")

two_col_row(doc, "Bezeichnung:",
    "Aufbewahrung von Lieferscheinen zur Rückverfolgbarkeit nach VO (EG) 178/2002")
two_col_row(doc, "Zweck:",
    "Lebensmittelrechtliche Rückverfolgbarkeit; Nachweis der Lieferkette")
two_col_row(doc, "Rechtsgrundlage:",
    "Art. 6 Abs. 1 lit. c DSGVO (rechtliche Verpflichtung)\n"
    "VO (EG) 178/2002 Art. 18 (Rückverfolgbarkeit)")
two_col_row(doc, "Kategorien betroffener Personen:",
    "Lieferanten (Florida Eis / Hinnis)")
two_col_row(doc, "Kategorien personenbezogener Daten:",
    "Name, Adresse, Lieferdaten des Lieferanten (auf Lieferscheinen)")
two_col_row(doc, "Empfänger:",
    "Keine Weitergabe; interne Aufbewahrung")
two_col_row(doc, "Speicherdauer:",
    "Mindestens 2 Jahre nach MHD des gelieferten Produkts")
two_col_row(doc, "Zuletzt aktualisiert:", HEUTE)

footer_line(doc)


# ═══════════════════════════════════════════════════════════════════════════════
#  SEITE 3: TOMs
# ═══════════════════════════════════════════════════════════════════════════════
s2 = doc.add_section(WD_SECTION.NEW_PAGE); set_a4(s2)

logo_header(doc, "Technische und organisatorische Maßnahmen (TOMs)", "Art. 32 DSGVO")
page_title(doc, "B", "Technische und organisatorische Maßnahmen (TOMs)",
           "Art. 32 DSGVO  ·  Bezug: Videoüberwachung Eisautomat-Standort")
gap(doc, 6)

cw = [Cm(6.5), Cm(2.2), Cm(8.7)]

section_bar(doc, "1. Zutrittskontrolle  (physischer Schutz)")
check_table(doc, [
    ["Kameragehäuse wetterfest und manipulationssicher montiert", "Ja", "Außenmontage in gesicherter Halterung"],
    ["Zugang zur Speichereinheit nur für Betreiber", "Ja", "Gerät verschlossen, Schlüssel bei Betreiber"],
    ["Kamerasichtfeld auf eigenes Grundstück / Automat begrenzt", "Ja", "Keine Erfassung öffentlicher Gehwege"],
], cw)

section_bar(doc, "2. Zugangskontrolle  (digitaler Schutz)")
check_table(doc, [
    ["Speichermedium / Aufzeichnungsgerät passwortgeschützt", "Ja", "Starkes Passwort, nur Betreiber bekannt"],
    ["Fernzugriff deaktiviert oder durch VPN/Passwort gesichert", "Ja", "Kein offener Fernzugriff"],
    ["Software / Firmware des Kamerasystems aktuell gehalten", "Ja", "Regelmäßige Updates werden eingespielt"],
], cw)

section_bar(doc, "3. Zugriffskontrolle  (Berechtigungen)")
check_table(doc, [
    ["Aufnahmen nur durch Betreiber einsehbar", "Ja", "Kein Zugriff für Dritte ohne konkreten Anlass"],
    ["Kein automatisiertes Auslesen / Analyse der Aufnahmen", "Ja", "Aufnahmen werden nicht automatisch ausgewertet"],
    ["Weitergabe nur bei dokumentiertem Vorfall an Behörden", "Ja", "Protokoll bei Herausgabe wird geführt"],
], cw)

section_bar(doc, "4. Verfügbarkeitskontrolle  (Datensicherheit)")
check_table(doc, [
    ["Automatische Überschreibung / Löschung nach 7 Tagen", "Ja", "Eingestellt am Aufzeichnungsgerät"],
    ["Schutz vor Datenverlust (Gerät gegen Diebstahl gesichert)", "Ja", "Gehäuse abschließbar"],
    ["Aufzeichnungsgerät vor Witterungseinflüssen geschützt", "Ja", "Wetterfestes Gehäuse"],
], cw)

section_bar(doc, "5. Trennbarkeit / Transparenz")
check_table(doc, [
    ["Hinweisschild sichtbar angebracht (Art. 13 DSGVO)", "Ja", "Aushang vor dem Kamerabereich"],
    ["Informationsblatt mit Betroffenenrechten ausgehängt", "Ja", "Aushang am Standort"],
    ["VVT geführt und aktuell gehalten", "Ja", "Dieses Dokument"],
], cw)

footer_line(doc)


# ═══════════════════════════════════════════════════════════════════════════════
#  SEITE 4: Löschkonzept
# ═══════════════════════════════════════════════════════════════════════════════
s3 = doc.add_section(WD_SECTION.NEW_PAGE); set_a4(s3)

logo_header(doc, "Löschkonzept", "Art. 5 Abs. 1 lit. e · Art. 17 DSGVO")
page_title(doc, "C", "Löschkonzept",
           "Art. 5 Abs. 1 lit. e DSGVO (Speicherbegrenzung)  ·  Art. 17 DSGVO (Recht auf Löschung)")
gap(doc, 6)

section_bar(doc, "Videoüberwachung – Löschfristen und Verfahren")

two_col_row(doc, "Datenart:", "Videoaufnahmen (Bilddaten / Bewegtbilder)")
two_col_row(doc, "Regelmäßige Löschfrist:", "7 Tage nach Aufnahme")
two_col_row(doc, "Löschmethode:", "Automatische Überschreibung durch das Aufzeichnungsgerät nach Ablauf von 7 Tagen")
two_col_row(doc, "Zuständigkeit:", "Marta und Florian Schmidl (Betreiber)")
two_col_row(doc, "Ausnahme – Vorfall:",
    "Bei dokumentiertem Schadensereignis (Vandalismus, Diebstahl) werden die betreffenden "
    "Aufnahmen gesichert und bis zur abschließenden Klärung des Vorfalls aufbewahrt. "
    "Nach Abschluss erfolgt unverzügliche Löschung.")
two_col_row(doc, "Ausnahme – Behördenanfrage:",
    "Fordert eine Strafverfolgungsbehörde Aufnahmen an, werden diese übergeben und "
    "anschließend lokal gelöscht. Die Herausgabe wird protokolliert.")
two_col_row(doc, "Nachweis / Kontrolle:",
    "Betreiber prüft monatlich, ob die automatische Überschreibung korrekt funktioniert. "
    "Abweichungen werden im Abweichungsprotokoll dokumentiert.")

gap(doc, 6)
section_bar(doc, "Lieferantendaten / Belege – Löschfristen")

two_col_row(doc, "Datenart:", "Lieferscheine, Belege (personenbezogene Lieferantendaten)")
two_col_row(doc, "Aufbewahrungsfrist:",
    "Mindestens 2 Jahre nach MHD des Produkts (VO (EG) 178/2002)\n"
    "Steuerrechtlich: 10 Jahre (§ 147 AO) – längere Frist gilt vorrangig")
two_col_row(doc, "Löschmethode:", "Physische Vernichtung (Schreddern) nach Ablauf der Aufbewahrungsfrist")
two_col_row(doc, "Zuständigkeit:", "Marta und Florian Schmidl")

gap(doc, 8)

# Unterschriftszeile
section_bar(doc, "Bestätigung und Aktualisierungsnachweis", bg=DUNKELBLAU)
gap(doc, 4)
sig_tbl = doc.add_table(rows=1, cols=2)
sig_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
sig_tbl.columns[0].width = Cm(8.7)
sig_tbl.columns[1].width = Cm(8.7)
for i, label in enumerate(["Datum / Unterschrift (Erstellt)", "Datum / Unterschrift (Geprüft)"]):
    c = sig_tbl.rows[0].cells[i]
    cell_bg(c, HELLGRAU); cell_border(c)
    p = c.paragraphs[0]
    p.paragraph_format.left_indent  = Cm(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(30)
    run(p, label, size=9, color=GRUEN)
    p2 = c.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.3)
    p2.paragraph_format.space_after = Pt(6)
    run(p2, "________________________________", size=10, color=SCHRIFT)

footer_line(doc)

doc.save(OUTPUT)
print("Gespeichert:", OUTPUT)
