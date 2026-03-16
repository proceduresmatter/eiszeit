"""
Erstellt ein DSGVO-konformes Hinweisschild (Videoüberwachung) nach Art. 13 DSGVO
für Eiszeit – Marta und Florian Schmidl
Vorlage: Landesbeauftragter für Datenschutz Niedersachsen
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Markenfarben Eiszeit ──────────────────────────────────────────────────────
BEIGE      = RGBColor(0xF0, 0xE4, 0xCB)   # Hintergrund
BLAU       = RGBColor(0x9B, 0xD4, 0xDB)   # Akzent
GRUEN      = RGBColor(0x6C, 0x94, 0x8C)   # Dunklerer Akzent
SCHRIFT    = RGBColor(0x41, 0x45, 0x45)   # Dunkelgrau
WEISS      = RGBColor(0xFF, 0xFF, 0xFF)
DUNKELBLAU = RGBColor(0x1A, 0x5F, 0x7A)   # Für Überschriften

# ── Pflichtangaben ────────────────────────────────────────────────────────────
DATEN = {
    "verantwortlicher": "Marta und Florian Schmidl",
    "adresse":          "Heiligenbruch 27, 27339 Riede",
    "email":            "eiszeit@mailbox.org",
    "dsb":              None,   # kein betrieblicher DSB
    "zweck":            "Vandalismusprävention, Schutz des Eigentums, Hausrecht",
    "rechtsgrundlage":  "Art. 6 Abs. 1 lit. f DSGVO (berechtigtes Interesse)",
    "interesse":        "Schutz des Eigentums und der Vertrauenskasse vor Vandalismus und Diebstahl",
    "speicherdauer":    "7 Tage, danach automatische Löschung",
    "info_ort":         "Aushang am Standort (ausführliches Informationsblatt)",
    "aufsicht":         "Die Landesbeauftragte für den Datenschutz Niedersachsen\n"
                        "Prinzenstr. 5, 30159 Hannover\n"
                        "Tel.: 0511 120 4500 · poststelle@lfd.niedersachsen.de",
}

LOGO_PATH   = "EiszeitLogo.png"
CAMERA_PATH = "kamera_piktogramm.png"
OUTPUT      = "DSGVO_Hinweisschild_Eiszeit.docx"


def set_cell_bg(cell, rgb: RGBColor):
    """Setzt die Hintergrundfarbe einer Tabellenzelle."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    tcPr.append(shd)


def set_cell_border(cell, sides=("top","bottom","left","right"), size=12, color="9BD4DB"):
    """Setzt Rahmen einer Zelle."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in sides:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def add_run(para, text, bold=False, italic=False, size=11,
            color=None, font="Calibri"):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def add_heading_row(table, text, col_span=2, bg=BLAU):
    """Fügt eine farbige Überschriftszeile über volle Breite ein."""
    row  = table.add_row()
    cell = row.cells[0]
    # Zellen verbinden
    if col_span == 2:
        cell.merge(row.cells[1])
    set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.3)
    add_run(p, text, bold=True, size=10, color=WEISS)
    return row


def add_content_row(table, label, value):
    """Fügt eine zweispaltige Inhaltszeile ein."""
    row = table.add_row()
    # linke Zelle (Label)
    lc = row.cells[0]
    set_cell_bg(lc, BEIGE)
    set_cell_border(lc, color="9BD4DB", size=6)
    pl = lc.paragraphs[0]
    pl.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pl.paragraph_format.left_indent  = Cm(0.3)
    pl.paragraph_format.space_before = Pt(3)
    pl.paragraph_format.space_after  = Pt(3)
    add_run(pl, label, bold=True, size=9.5, color=GRUEN)
    # rechte Zelle (Wert)
    rc = row.cells[1]
    set_cell_bg(rc, WEISS)
    set_cell_border(rc, color="9BD4DB", size=6)
    pr = rc.paragraphs[0]
    pr.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pr.paragraph_format.left_indent  = Cm(0.3)
    pr.paragraph_format.space_before = Pt(3)
    pr.paragraph_format.space_after  = Pt(3)
    add_run(pr, value, size=9.5, color=SCHRIFT)
    return row


# ── Dokument aufbauen ─────────────────────────────────────────────────────────
doc = Document()

# Seitenformat DIN A4, schmale Ränder
section = doc.sections[0]
section.page_width   = Cm(21)
section.page_height  = Cm(29.7)
section.top_margin    = Cm(1.5)
section.bottom_margin = Cm(1.5)
section.left_margin   = Cm(1.8)
section.right_margin  = Cm(1.8)

# ── Kopfzeile: Logo | Kamera-Piktogramm | Titel ──────────────────────────────
header_tbl = doc.add_table(rows=1, cols=3)
header_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
header_tbl.columns[0].width = Cm(3.2)   # Eiszeit-Logo
header_tbl.columns[1].width = Cm(4.2)   # Kamera-Piktogramm
header_tbl.columns[2].width = Cm(10.0)  # Titeltext

logo_cell   = header_tbl.rows[0].cells[0]
cam_cell    = header_tbl.rows[0].cells[1]
title_cell  = header_tbl.rows[0].cells[2]

for c in (logo_cell, cam_cell, title_cell):
    set_cell_bg(c, BEIGE)

# Eiszeit-Logo
lp = logo_cell.paragraphs[0]
lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
lp.paragraph_format.space_before = Pt(8)
lp.paragraph_format.space_after  = Pt(8)
try:
    lp.add_run().add_picture(LOGO_PATH, width=Cm(2.8))
except Exception:
    add_run(lp, "EISZEIT", bold=True, size=14, color=DUNKELBLAU)

# Kamera-Piktogramm (gross, zentriert)
cp = cam_cell.paragraphs[0]
cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
cp.paragraph_format.space_before = Pt(6)
cp.paragraph_format.space_after  = Pt(6)
try:
    cp.add_run().add_picture(CAMERA_PATH, width=Cm(3.6))
except Exception:
    add_run(cp, "📹", size=40)

# Titeltext
tp = title_cell.paragraphs[0]
tp.alignment = WD_ALIGN_PARAGRAPH.LEFT
tp.paragraph_format.space_before = Pt(10)
tp.paragraph_format.left_indent  = Cm(0.4)
add_run(tp, "Achtung!", bold=True, size=28, color=DUNKELBLAU)

tp2 = title_cell.add_paragraph()
tp2.alignment = WD_ALIGN_PARAGRAPH.LEFT
tp2.paragraph_format.left_indent  = Cm(0.4)
tp2.paragraph_format.space_before = Pt(2)
add_run(tp2, "Videoüberwachung!", bold=True, size=18, color=SCHRIFT)

tp3 = title_cell.add_paragraph()
tp3.alignment = WD_ALIGN_PARAGRAPH.LEFT
tp3.paragraph_format.left_indent  = Cm(0.4)
tp3.paragraph_format.space_before = Pt(6)
add_run(tp3, "Dieser Bereich wird durch\neine Videokamera überwacht.",
        size=9, italic=True, color=SCHRIFT)

doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ── Pflichtangaben-Tabelle ────────────────────────────────────────────────────
tbl = doc.add_table(rows=0, cols=2)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.columns[0].width = Cm(5.5)
tbl.columns[1].width = Cm(11.9)

# 1 · Verantwortlicher
add_heading_row(tbl, "📋  Verantwortlicher (Art. 13 Abs. 1 lit. a DSGVO)", bg=GRUEN)
add_content_row(tbl, "Name:",    DATEN["verantwortlicher"])
add_content_row(tbl, "Adresse:", DATEN["adresse"])
add_content_row(tbl, "E-Mail:",  DATEN["email"])

# 2 · DSB
add_heading_row(tbl, "📋  Datenschutzbeauftragter (Art. 13 Abs. 1 lit. b DSGVO)", bg=GRUEN)
add_content_row(tbl, "Hinweis:",
    "Es wurde kein betrieblicher Datenschutzbeauftragter bestellt.")

# 3 · Zweck & Rechtsgrundlage
add_heading_row(tbl, "🎯  Zweck und Rechtsgrundlage (Art. 13 Abs. 1 lit. c DSGVO)", bg=GRUEN)
add_content_row(tbl, "Zweck:",          DATEN["zweck"])
add_content_row(tbl, "Rechtsgrundlage:", DATEN["rechtsgrundlage"])
add_content_row(tbl, "Berechtigtes\nInteresse:", DATEN["interesse"])

# 4 · Speicherdauer
add_heading_row(tbl, "🕐  Speicherdauer (Art. 13 Abs. 2 lit. a DSGVO)", bg=GRUEN)
add_content_row(tbl, "Dauer:", DATEN["speicherdauer"])

# 5 · Weiterführende Infos
add_heading_row(tbl, "ℹ️  Weitere Informationen", bg=DUNKELBLAU)
add_content_row(tbl, "Informationsblatt:",
    "Das vollständige Informationsblatt mit Ihren Betroffenenrechten "
    "(Auskunft, Widerspruch, Löschung, Beschwerde) finden Sie als Aushang am Standort.")
add_content_row(tbl, "Aufsichtsbehörde:", DATEN["aufsicht"])

# ── Fußzeile ──────────────────────────────────────────────────────────────────
doc.add_paragraph()
fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.paragraph_format.space_before = Pt(6)
add_run(fp,
    "Erstellt gemäß Art. 12–13 DSGVO · Vorlage: LfD Niedersachsen",
    size=7.5, italic=True, color=RGBColor(0xAA, 0xAA, 0xAA))

doc.save(OUTPUT)
print(f"✓ Gespeichert: {OUTPUT}")
