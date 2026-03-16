from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def create_eiszeit_qm():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    # ===== HEADER =====
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run('EISZEIT')
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(70, 70, 70)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Kontrollkonzept & Qualitätsdokumentation')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    # ===== BETRIEBSDATEN =====
    info = doc.add_paragraph()
    info.add_run('Standort: ').bold = True
    info.add_run('SB-Eisautomat Eiszeit\n')
    info.add_run('Betreiber: ').bold = True
    info.add_run('[Name/Firma]\n')
    info.add_run('Dokumentationszeitraum: ').bold = True
    info.add_run('______________ bis ______________')

    doc.add_paragraph()

    # ===== GRUNDSATZ =====
    grundsatz_table = doc.add_table(rows=1, cols=1)
    grundsatz_table.style = 'Table Grid'
    cell = grundsatz_table.rows[0].cells[0]
    set_cell_shading(cell, 'FFF9E6')
    para = cell.paragraphs[0]
    para.add_run('Dokumentationsprinzip: ').bold = True
    para.add_run('Alle Kontrollvorgaben sind in diesem Dokument definiert. ')
    para.add_run('Solange alles im Zielbereich liegt, ist keine Eintragung erforderlich. ').italic = True
    para.add_run('Nur Abweichungen und Korrekturmaßnahmen werden dokumentiert.')

    doc.add_paragraph()

    # ===== 1. KONTROLLKONZEPT =====
    doc.add_heading('1. Kontrollkonzept', level=1)

    # --- 1.1 Temperaturkontrolle ---
    doc.add_heading('1.1 Temperaturkontrolle', level=2)

    table_temp = doc.add_table(rows=3, cols=4)
    table_temp.style = 'Table Grid'

    headers = ['Bereich', 'Zielwert', 'Kontrolle', 'Dokumentation']
    for i, h in enumerate(headers):
        cell = table_temp.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'B8D4E8')

    temp_data = [
        ['Wareneingang', '≤ -18 °C', 'Bei jeder Anlieferung mit Thermometer prüfen', 'Nur bei Abweichung'],
        ['Automat / Lager', '≤ -18 °C', '1x täglich (Display/Thermometer) oder Datenlogger', 'Nur bei Abweichung']
    ]
    for row_idx, row_data in enumerate(temp_data):
        for col_idx, text in enumerate(row_data):
            table_temp.rows[row_idx + 1].cells[col_idx].text = text

    doc.add_paragraph()

    # --- 1.2 Reinigungskonzept ---
    doc.add_heading('1.2 Reinigungskonzept', level=2)

    table_rein = doc.add_table(rows=4, cols=4)
    table_rein.style = 'Table Grid'

    headers_rein = ['Bereich', 'Häufigkeit', 'Durchführung', 'Reinigungsmittel']
    for i, h in enumerate(headers_rein):
        cell = table_rein.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'B8E8D4')

    rein_data = [
        ['Außenflächen Automat', '1x täglich', 'Abwischen mit Tuch & Oberflächenreiniger', 'Lebensmitteltauglich'],
        ['Innenraum (Fächer)', '1x wöchentlich', 'Entleeren, Fächer reinigen', 'Lebensmitteltauglich'],
        ['Außenbereich', '1x täglich', 'Mülleimer leeren, Sichtkontrolle Umfeld', '–']
    ]
    for row_idx, row_data in enumerate(rein_data):
        for col_idx, text in enumerate(row_data):
            table_rein.rows[row_idx + 1].cells[col_idx].text = text

    hinweis_rein = doc.add_paragraph()
    hinweis_rein.add_run('→ Reinigung gilt als durchgeführt, solange keine Abweichung dokumentiert ist.').italic = True

    doc.add_paragraph()

    # --- 1.3 Sonstige Maßnahmen ---
    doc.add_heading('1.3 Sonstige Maßnahmen', level=2)

    table_sonst = doc.add_table(rows=4, cols=3)
    table_sonst.style = 'Table Grid'

    headers_sonst = ['Maßnahme', 'Häufigkeit', 'Hinweis']
    for i, h in enumerate(headers_sonst):
        cell = table_sonst.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D4B8E8')

    sonst_data = [
        ['Wartung Automat', '1x pro Quartal', 'Durch Techniker, Dokumentation unter Punkt 5'],
        ['Schädlingskontrolle', '1x wöchentlich', 'Sichtprüfung im Umkreis des Automaten'],
        ['Datenlogger auslesen', 'Nach Bedarf', 'Falls automatische Temperaturüberwachung vorhanden']
    ]
    for row_idx, row_data in enumerate(sonst_data):
        for col_idx, text in enumerate(row_data):
            table_sonst.rows[row_idx + 1].cells[col_idx].text = text

    doc.add_paragraph()

    # --- 1.4 Rückverfolgbarkeit ---
    doc.add_heading('1.4 Rückverfolgbarkeit', level=2)

    rueck = doc.add_paragraph()
    rueck.add_run('Gemäß VO (EG) 178/2002 ist die Rückverfolgbarkeit sicherzustellen:\n\n')
    rueck.add_run('Lieferant: ').bold = True
    rueck.add_run('[Hauptlieferant eintragen]\n')
    rueck.add_run('Produkte: ').bold = True
    rueck.add_run('Vorverpacktes Speiseeis\n')
    rueck.add_run('Dokumentation: ').bold = True
    rueck.add_run('Lieferscheine werden aufbewahrt (mind. 2 Jahre nach MHD)')

    doc.add_paragraph()

    # --- 1.5 Personalschulung ---
    doc.add_heading('1.5 Personalschulung', level=2)

    schulung = doc.add_paragraph()
    schulung.add_run('Alle Personen, die mit Lebensmitteln umgehen, müssen geschult sein:\n\n')

    schulung_items = [
        'Erstbelehrung nach § 43 IfSG (Gesundheitsamt) – falls erforderlich',
        'HACCP-Grundlagen und Lebensmittelhygiene',
        'Folgebelehrungen: jährlich (intern dokumentiert)'
    ]
    for item in schulung_items:
        doc.add_paragraph(item, style='List Bullet')

    hinweis_ifsg = doc.add_paragraph()
    hinweis_ifsg.add_run('Hinweis: ').bold = True
    hinweis_ifsg.add_run('Bei ausschließlich vorverpackter Ware ist die IfSG-Belehrung nach Rücksprache '
                         'mit dem Landkreis Verden nicht erforderlich.').italic = True

    doc.add_paragraph()

    # --- 1.6 Rechtliche Grundlagen ---
    doc.add_heading('1.6 Rechtliche Grundlagen', level=2)

    rechtlich = doc.add_paragraph()
    rechtlich.add_run('Anwendbare Vorschriften:\n\n')

    vorschriften = [
        'VO (EG) 852/2004 – Lebensmittelhygiene',
        'VO (EG) 178/2002 – Allgemeines Lebensmittelrecht, Rückverfolgbarkeit',
        'LMHV – Lebensmittelhygiene-Verordnung (Deutschland)',
        'IfSG §§ 42, 43 – Infektionsschutzgesetz (bei Bedarf)'
    ]
    for v in vorschriften:
        doc.add_paragraph(v, style='List Bullet')

    doc.add_paragraph()

    # --- 1.7 Produktrückruf ---
    doc.add_heading('1.7 Notfallverfahren / Produktrückruf', level=2)

    notfall = doc.add_paragraph()
    notfall.add_run('Bei Verdacht auf gesundheitsgefährdende Produkte:\n\n')

    notfall_steps = [
        'Betroffene Ware sofort aus dem Verkauf nehmen',
        'Ware separieren und kennzeichnen („Nicht zum Verkauf")',
        'Lieferant und ggf. Lebensmittelüberwachung informieren',
        'Vorfall im Abweichungsprotokoll dokumentieren',
        'Kaufbelege/Lieferscheine für Rückverfolgung bereithalten'
    ]
    for i, step in enumerate(notfall_steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    doc.add_page_break()

    # ===== 2. EISZEIT UNTERWEGS =====
    doc.add_heading('2. Eiszeit unterwegs – Mobiler Verkauf', level=1)

    unterwegs_intro = doc.add_paragraph()
    unterwegs_intro.add_run('Besondere Anforderungen bei Verkauf auf Festen, Märkten und Veranstaltungen:').italic = True

    doc.add_paragraph()

    # --- 2.1 Vor dem Einsatz ---
    doc.add_heading('2.1 Vorbereitung (vor dem Einsatz)', level=2)

    table_vor = doc.add_table(rows=7, cols=2)
    table_vor.style = 'Table Grid'

    headers_vor = ['Prüfpunkt', 'Erledigt']
    for i, h in enumerate(headers_vor):
        cell = table_vor.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'FFD4A8')  # Orange pastel

    vor_data = [
        ['Genehmigung/Anmeldung beim Veranstalter oder Ordnungsamt', '☐'],
        ['Stromversorgung am Standort geklärt (Kühlung!)', '☐'],
        ['Kühlbox / mobile Kühlung funktionsfähig', '☐'],
        ['Thermometer eingepackt', '☐'],
        ['Reinigungsmittel & Tücher dabei', '☐'],
        ['Müllbeutel / Entsorgungsmöglichkeit', '☐']
    ]
    for row_idx, row_data in enumerate(vor_data):
        for col_idx, text in enumerate(row_data):
            table_vor.rows[row_idx + 1].cells[col_idx].text = text

    doc.add_paragraph()

    # --- 2.2 Kühlkette ---
    doc.add_heading('2.2 Kühlkette beim Transport', level=2)

    kuehl_box = doc.add_table(rows=1, cols=1)
    kuehl_box.style = 'Table Grid'
    cell = kuehl_box.rows[0].cells[0]
    set_cell_shading(cell, 'FFE6E6')
    para = cell.paragraphs[0]
    para.add_run('WICHTIG: ').bold = True
    para.add_run('Die Kühlkette darf nicht unterbrochen werden!\n')
    para.add_run('Speiseeis muss während des gesamten Transports bei ≤ -18 °C gehalten werden.')

    doc.add_paragraph()

    kuehl_hinweise = [
        'Transport in geeigneter Kühlbox oder Tiefkühltasche mit Kühlelementen',
        'Transportzeit so kurz wie möglich halten',
        'Temperatur bei Ankunft prüfen und dokumentieren',
        'Bei Unterbrechung der Kühlkette (> -15 °C): Ware NICHT mehr verkaufen',
        'Angetautes Eis darf NICHT wieder eingefroren werden'
    ]
    for h in kuehl_hinweise:
        doc.add_paragraph(h, style='List Bullet')

    doc.add_paragraph()

    # --- 2.3 Vor Ort ---
    doc.add_heading('2.3 Vor Ort auf der Veranstaltung', level=2)

    table_ort = doc.add_table(rows=6, cols=3)
    table_ort.style = 'Table Grid'

    headers_ort = ['Bereich', 'Anforderung', 'Hinweis']
    for i, h in enumerate(headers_ort):
        cell = table_ort.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'FFD4A8')

    ort_data = [
        ['Standort', 'Sauber, trocken, schattig wenn möglich', 'Direkte Sonne vermeiden'],
        ['Kühlung', 'Sofort an Strom anschließen', 'Temperatur regelmäßig prüfen'],
        ['Hygiene', 'Hände waschen / desinfizieren', 'Handdesinfektion mitnehmen'],
        ['Bezahlung', 'Geld und Ware getrennt handhaben', 'Erst Geld, dann Eis ausgeben'],
        ['Abfall', 'Mülleimer bereitstellen', 'Regelmäßig leeren']
    ]
    for row_idx, row_data in enumerate(ort_data):
        for col_idx, text in enumerate(row_data):
            table_ort.rows[row_idx + 1].cells[col_idx].text = text

    doc.add_paragraph()

    # --- 2.4 Notfall unterwegs ---
    doc.add_heading('2.4 Notfall: Kühlung fällt aus', level=2)

    notfall_unterwegs = doc.add_paragraph()
    notfall_unterwegs.add_run('Wenn die Kühlung unterwegs oder vor Ort ausfällt:\n\n')

    notfall_steps_uw = [
        'Verkauf sofort stoppen',
        'Temperatur der Ware messen',
        'Wenn Eis angetaut (> -15 °C über längere Zeit): Ware entsorgen',
        'Ersatzkühlung organisieren oder Veranstaltung abbrechen',
        'Vorfall dokumentieren (Datum, Uhrzeit, Maßnahme)'
    ]
    for i, step in enumerate(notfall_steps_uw, 1):
        doc.add_paragraph(f'{i}. {step}')

    doc.add_paragraph()

    # --- 2.5 Checkliste Eiszeit unterwegs ---
    doc.add_heading('2.5 Packliste Eiszeit unterwegs', level=2)

    pack_items = [
        ('Ware', 'Eis in ausreichender Menge, vorverpackt'),
        ('Kühlung', 'Kühlbox, Kühlelemente, ggf. Verlängerungskabel'),
        ('Thermometer', 'Zur Temperaturkontrolle'),
        ('Hygiene', 'Handdesinfektion, Reinigungstücher, Papiertücher'),
        ('Entsorgung', 'Müllbeutel'),
        ('Geld', 'Wechselgeld, Kasse'),
        ('Dokumentation', 'Dieses Dokument, Stift'),
        ('Kontakt', 'Telefonnummer Techniker / Lieferant')
    ]

    table_pack = doc.add_table(rows=len(pack_items)+1, cols=3)
    table_pack.style = 'Table Grid'

    pack_headers = ['Kategorie', 'Was', '✓']
    for i, h in enumerate(pack_headers):
        cell = table_pack.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'FFD4A8')

    for row_idx, (kat, was) in enumerate(pack_items):
        table_pack.rows[row_idx + 1].cells[0].text = kat
        table_pack.rows[row_idx + 1].cells[1].text = was
        table_pack.rows[row_idx + 1].cells[2].text = '☐'

    doc.add_page_break()

    # ===== 3. ABWEICHUNGSPROTOKOLL =====
    doc.add_heading('3. Abweichungsprotokoll', level=1)

    erklaerung = doc.add_paragraph()
    erklaerung.add_run('Hier werden ').italic = True
    erklaerung.add_run('nur Abweichungen').bold = True
    erklaerung.add_run(' dokumentiert: Temperatur außerhalb Zielbereich, Verschmutzung, Beschädigung, '
                       'Schädlingshinweis, Lieferprobleme, technische Störungen, Vorfälle unterwegs, etc.').italic = True

    doc.add_paragraph()

    table_abw = doc.add_table(rows=22, cols=5)
    table_abw.style = 'Table Grid'

    headers_abw = ['Datum', 'Bereich', 'Feststellung', 'Maßnahme', 'Kürzel']
    for i, h in enumerate(headers_abw):
        cell = table_abw.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'F4B8C5')

    beispiele = [
        ['15.06.', 'Temperatur', '-15°C (zu warm)', 'Techniker gerufen, Ware kontrolliert', 'MS'],
        ['22.06.', 'Wareneingang', 'Karton beschädigt', 'Ware geprüft, i.O., Lieferant informiert', 'MS'],
    ]
    for row_idx, row_data in enumerate(beispiele):
        for col_idx, text in enumerate(row_data):
            cell = table_abw.rows[row_idx + 1].cells[col_idx]
            cell.text = text
            set_cell_shading(cell, 'FFF0F3')
            for para in cell.paragraphs:
                if para.runs:
                    para.runs[0].italic = True

    for row_idx in range(3, 22):
        for col_idx in range(5):
            table_abw.rows[row_idx].cells[col_idx].text = ''

    doc.add_paragraph()

    # ===== 4. WARENEINGANG =====
    doc.add_heading('4. Wareneingang (nur bei Abweichung)', level=1)

    waren_info = doc.add_paragraph()
    waren_info.add_run('Zielwerte: ').bold = True
    waren_info.add_run('Temperatur ≤ -18 °C | Verpackung unbeschädigt | MHD ausreichend\n')
    waren_info.add_run('→ Bei Lieferung ohne Beanstandung: keine Eintragung erforderlich').italic = True

    doc.add_paragraph()

    table_ware = doc.add_table(rows=10, cols=5)
    table_ware.style = 'Table Grid'

    headers_ware = ['Datum', 'Lieferant', 'Temp. °C', 'Feststellung', 'Maßnahme/Kürzel']
    for i, h in enumerate(headers_ware):
        cell = table_ware.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'FFF3B8')

    for row_idx in range(1, 10):
        for col_idx in range(5):
            table_ware.rows[row_idx].cells[col_idx].text = ''

    doc.add_paragraph()

    # ===== 5. WARTUNG & KONTROLLEN =====
    doc.add_heading('5. Wartung & Sonderkontrollen', level=1)

    doc.add_paragraph('Quartalswartung, Techniker-Einsätze, Behördenkontrollen, Schulungen, etc.')

    table_wart = doc.add_table(rows=10, cols=4)
    table_wart.style = 'Table Grid'

    headers_wart = ['Datum', 'Art der Maßnahme', 'Durchgeführt von', 'Ergebnis/Bemerkung']
    for i, h in enumerate(headers_wart):
        cell = table_wart.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D4B8E8')

    for row_idx in range(1, 10):
        for col_idx in range(4):
            table_wart.rows[row_idx].cells[col_idx].text = ''

    doc.add_paragraph()

    # ===== 6. EINSÄTZE UNTERWEGS =====
    doc.add_heading('6. Einsätze Eiszeit unterwegs', level=1)

    doc.add_paragraph('Dokumentation aller mobilen Einsätze auf Festen, Märkten, Veranstaltungen.')

    table_unterwegs = doc.add_table(rows=10, cols=5)
    table_unterwegs.style = 'Table Grid'

    headers_uw = ['Datum', 'Veranstaltung/Ort', 'Temp. Ankunft', 'Besonderheiten', 'Kürzel']
    for i, h in enumerate(headers_uw):
        cell = table_unterwegs.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'FFD4A8')

    for row_idx in range(1, 10):
        for col_idx in range(5):
            table_unterwegs.rows[row_idx].cells[col_idx].text = ''

    doc.add_paragraph()
    doc.add_paragraph()

    # ===== ABSCHLUSS =====
    abschluss = doc.add_paragraph()
    abschluss.add_run('Dokumentation abgeschlossen am: ').bold = True
    abschluss.add_run('______________\n\n')
    abschluss.add_run('Unterschrift: ').bold = True
    abschluss.add_run('______________________________')

    doc.add_paragraph()
    doc.add_paragraph()

    # Footer
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run('EISZEIT').bold = True
    footer.add_run(' | Kontrollkonzept & Qualitätsdokumentation | Version 3.0')

    # Save
    output = r'C:\Users\localUser\Desktop\Claude\Eiszeit\Eiszeit_Kontrollkonzept_v3.0.docx'
    doc.save(output)
    print(f'Document saved: {output}')
    return output

if __name__ == '__main__':
    create_eiszeit_qm()
