"""
Erstellt das vollständige Informationsblatt (Stufe 2, DIN A3) nach Art. 13 DSGVO
für Eiszeit – Marta und Florian Schmidl
Vorlage: Landesbeauftragter für Datenschutz Niedersachsen
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Farben ────────────────────────────────────────────────────────────────────
BEIGE      = RGBColor(0xF0, 0xE4, 0xCB)
BLAU       = RGBColor(0x9B, 0xD4, 0xDB)
GRUEN      = RGBColor(0x6C, 0x94, 0x8C)
SCHRIFT    = RGBColor(0x41, 0x45, 0x45)
WEISS      = RGBColor(0xFF, 0xFF, 0xFF)
DUNKELBLAU = RGBColor(0x1A, 0x5F, 0x7A)
HELLGRAU   = RGBColor(0xF5, 0xF5, 0xF5)

DATEN = {
    "verantwortlicher": "Marta und Florian Schmidl",
    "adresse":          "Heiligenbruch 27, 27339 Riede",
    "email":            "eiszeit@mailbox.org",
    "zweck":            "Vandalismusprävention, Schutz des Eigentums, Hausrecht",
    "rechtsgrundlage":  "Art. 6 Abs. 1 lit. f DSGVO (berechtigtes Interesse)",
    "interesse":        "Schutz des Eigentums und der Vertrauenskasse vor Vandalismus und Diebstahl",
    "speicherdauer":    "Die Aufnahmen werden automatisch nach 7 Tagen gelöscht, sofern kein konkreter Vorfall dokumentiert werden muss.",
    "empfaenger":       "Die Aufnahmen werden grundsätzlich nicht an Dritte weitergegeben. Im Fall eines dokumentierten Vorfalls können Aufnahmen an Strafverfolgungsbehörden übermittelt werden.",
    "drittland":        "Eine Übermittlung personenbezogener Daten an Drittländer oder internationale Organisationen findet nicht statt.",
    "aufsicht_name":    "Die Landesbeauftragte für den Datenschutz Niedersachsen",
    "aufsicht_adresse": "Prinzenstr. 5, 30159 Hannover",
    "aufsicht_tel":     "0511 120 4500",
    "aufsicht_email":   "poststelle@lfd.niedersachsen.de",
    "aufsicht_web":     "www.lfd.niedersachsen.de",
}

LOGO_PATH   = "EiszeitLogo.png"
CAMERA_PATH = "kamera_piktogramm.png"
OUTPUT      = "DSGVO_Informationsblatt_Eiszeit.docx"


def set_cell_bg(cell, rgb):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    tcPr.append(shd)


def set_cell_border(cell, color="9BD4DB", size=8):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "bottom", "left", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def p_run(para, text, bold=False, italic=False, size=10.5,
          color=None, font="Calibri"):
    run = para.add_run(text)
    run.bold        = bold
    run.italic      = italic
    run.font.name   = font
    run.font.size   = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def section_heading(doc, icon, title):
    """Farbiger Abschnittsbalken."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.columns[0].width = Cm(27)
    c = tbl.rows[0].cells[0]
    set_cell_bg(c, GRUEN)
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent  = Cm(0.4)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(5)
    p_run(p, f"{icon}  {title}", bold=True, size=12, color=WEISS)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def info_block(doc, label, text, label_color=GRUEN):
    """Zweispaltige Info-Zeile."""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.columns[0].width = Cm(6)
    tbl.columns[1].width = Cm(21)

    lc = tbl.rows[0].cells[0]
    rc = tbl.rows[0].cells[1]
    set_cell_bg(lc, BEIGE)
    set_cell_bg(rc, WEISS)
    set_cell_border(lc)
    set_cell_border(rc)

    pl = lc.paragraphs[0]
    pl.paragraph_format.left_indent  = Cm(0.3)
    pl.paragraph_format.space_before = Pt(4)
    pl.paragraph_format.space_after  = Pt(4)
    p_run(pl, label, bold=True, size=10, color=label_color)

    pr = rc.paragraphs[0]
    pr.paragraph_format.left_indent  = Cm(0.3)
    pr.paragraph_format.space_before = Pt(4)
    pr.paragraph_format.space_after  = Pt(4)
    p_run(pr, text, size=10, color=SCHRIFT)

    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(1)


def rights_block(doc, art_nr, titel, text):
    """Box für ein Betroffenenrecht."""
    tbl = doc.add_table(rows=2, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.columns[0].width = Cm(27)

    # Titelzeile
    hc = tbl.rows[0].cells[0]
    set_cell_bg(hc, BLAU)
    set_cell_border(hc, color="6C948C")
    ph = hc.paragraphs[0]
    ph.paragraph_format.left_indent  = Cm(0.4)
    ph.paragraph_format.space_before = Pt(4)
    ph.paragraph_format.space_after  = Pt(4)
    p_run(ph, f"{art_nr}  |  {titel}", bold=True, size=10.5, color=DUNKELBLAU)

    # Textzeile
    tc2 = tbl.rows[1].cells[0]
    set_cell_bg(tc2, HELLGRAU)
    set_cell_border(tc2, color="9BD4DB")
    pt = tc2.paragraphs[0]
    pt.paragraph_format.left_indent  = Cm(0.4)
    pt.paragraph_format.right_indent = Cm(0.4)
    pt.paragraph_format.space_before = Pt(4)
    pt.paragraph_format.space_after  = Pt(5)
    p_run(pt, text, size=9.5, color=SCHRIFT)

    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(2)


# ── Dokument ──────────────────────────────────────────────────────────────────
doc = Document()

section = doc.sections[0]
# DIN A3 Hochformat: 297 × 420 mm
section.page_width    = Cm(29.7)
section.page_height   = Cm(42.0)
section.top_margin    = Cm(1.8)
section.bottom_margin = Cm(1.8)
section.left_margin   = Cm(1.5)
section.right_margin  = Cm(1.5)

# ── Kopf ──────────────────────────────────────────────────────────────────────
head_tbl = doc.add_table(rows=1, cols=3)
head_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
head_tbl.columns[0].width = Cm(4.0)   # Eiszeit-Logo
head_tbl.columns[1].width = Cm(5.0)   # Kamera-Piktogramm
head_tbl.columns[2].width = Cm(18.0)  # Titeltext

logo_cell  = head_tbl.rows[0].cells[0]
cam_cell   = head_tbl.rows[0].cells[1]
title_cell = head_tbl.rows[0].cells[2]

for c in (logo_cell, cam_cell, title_cell):
    set_cell_bg(c, BEIGE)
    set_cell_border(c, color="9BD4DB", size=4)

# Eiszeit-Logo
lp = logo_cell.paragraphs[0]
lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
lp.paragraph_format.space_before = Pt(8)
try:
    lp.add_run().add_picture(LOGO_PATH, width=Cm(3.5))
except Exception:
    p_run(lp, "EISZEIT", bold=True, size=20, color=DUNKELBLAU)

# Kamera-Piktogramm
cp = cam_cell.paragraphs[0]
cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
cp.paragraph_format.space_before = Pt(6)
cp.paragraph_format.space_after  = Pt(6)
try:
    cp.add_run().add_picture(CAMERA_PATH, width=Cm(4.2))
except Exception:
    p_run(cp, "📹", size=48)

# Titeltext
tp = title_cell.paragraphs[0]
tp.alignment = WD_ALIGN_PARAGRAPH.LEFT
tp.paragraph_format.left_indent  = Cm(0.6)
tp.paragraph_format.space_before = Pt(10)
p_run(tp, "Datenschutzinformationen", bold=True, size=22, color=DUNKELBLAU)

tp2 = title_cell.add_paragraph()
tp2.alignment = WD_ALIGN_PARAGRAPH.LEFT
tp2.paragraph_format.left_indent = Cm(0.6)
p_run(tp2, "Vollständiges Informationsblatt zur Videoüberwachung",
      size=13, color=SCHRIFT)

tp3 = title_cell.add_paragraph()
tp3.alignment = WD_ALIGN_PARAGRAPH.LEFT
tp3.paragraph_format.left_indent  = Cm(0.6)
tp3.paragraph_format.space_before = Pt(4)
p_run(tp3,
      "gemäß Art. 13 der Datenschutz-Grundverordnung (DSGVO) · Standort: Heiligenbruch 27, 27339 Riede",
      size=9, italic=True, color=SCHRIFT)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── Abschnitt 1: Verantwortlicher ─────────────────────────────────────────────
section_heading(doc, "1.", "Verantwortlicher (Art. 13 Abs. 1 lit. a DSGVO)")
info_block(doc, "Name:",    DATEN["verantwortlicher"])
info_block(doc, "Adresse:", DATEN["adresse"])
info_block(doc, "E-Mail:",  DATEN["email"])

# ── Abschnitt 2: DSB ──────────────────────────────────────────────────────────
section_heading(doc, "2.", "Datenschutzbeauftragter (Art. 13 Abs. 1 lit. b DSGVO)")
info_block(doc, "Hinweis:",
    "Es wurde kein betrieblicher Datenschutzbeauftragter bestellt, "
    "da die Voraussetzungen gemäß Art. 37 DSGVO nicht vorliegen.")

# ── Abschnitt 3: Zweck & Rechtsgrundlage ──────────────────────────────────────
section_heading(doc, "3.", "Zwecke und Rechtsgrundlage (Art. 13 Abs. 1 lit. c–d DSGVO)")
info_block(doc, "Zweck:",             DATEN["zweck"])
info_block(doc, "Rechtsgrundlage:",   DATEN["rechtsgrundlage"])
info_block(doc, "Berechtigtes\nInteresse:", DATEN["interesse"])

# ── Abschnitt 4: Speicherdauer ────────────────────────────────────────────────
section_heading(doc, "4.", "Speicherdauer (Art. 13 Abs. 2 lit. a DSGVO)")
info_block(doc, "Dauer:", DATEN["speicherdauer"])

# ── Abschnitt 5: Empfänger ────────────────────────────────────────────────────
section_heading(doc, "5.", "Empfänger der Daten (Art. 13 Abs. 1 lit. e DSGVO)")
info_block(doc, "Empfänger:",    DATEN["empfaenger"])
info_block(doc, "Drittländer:",  DATEN["drittland"])

# ── Abschnitt 6: Betroffenenrechte ────────────────────────────────────────────
section_heading(doc, "6.", "Ihre Rechte als betroffene Person")

rights_block(doc,
    "Art. 15 DSGVO", "Recht auf Auskunft",
    "Sie haben das Recht, von uns eine Bestätigung darüber zu verlangen, ob Sie betreffende "
    "personenbezogene Daten verarbeitet werden. Ist dies der Fall, haben Sie ein Recht auf "
    "Auskunft über diese Daten sowie die in Art. 15 DSGVO aufgeführten Informationen.")

rights_block(doc,
    "Art. 16 DSGVO", "Recht auf Berichtigung",
    "Sie haben das Recht, von uns unverzüglich die Berichtigung Sie betreffender unrichtiger "
    "personenbezogener Daten sowie ggf. die Vervollständigung unvollständiger Daten zu verlangen.")

rights_block(doc,
    "Art. 17 DSGVO", "Recht auf Löschung",
    "Sie haben das Recht, von uns zu verlangen, dass Sie betreffende personenbezogene Daten "
    "unverzüglich gelöscht werden, sofern einer der in Art. 17 DSGVO genannten Gründe zutrifft – "
    "z. B. wenn die Daten für die verfolgten Zwecke nicht mehr benötigt werden.")

rights_block(doc,
    "Art. 18 DSGVO", "Recht auf Einschränkung der Verarbeitung",
    "Sie haben das Recht, von uns die Einschränkung der Verarbeitung zu verlangen, wenn eine der "
    "in Art. 18 DSGVO aufgeführten Voraussetzungen gegeben ist – z. B. wenn Sie Widerspruch gegen "
    "die Verarbeitung eingelegt haben, für die Dauer der Prüfung durch uns.")

rights_block(doc,
    "Art. 21 DSGVO", "Recht auf Widerspruch",
    "Sie haben das Recht, aus Gründen, die sich aus Ihrer besonderen Situation ergeben, jederzeit "
    "Widerspruch gegen die Verarbeitung Sie betreffender personenbezogener Daten einzulegen. "
    "Wir verarbeiten die Daten dann nicht mehr, es sei denn, wir können zwingende schutzwürdige "
    "Gründe nachweisen, die Ihre Interessen überwiegen.")

rights_block(doc,
    "Art. 77 DSGVO", "Recht auf Beschwerde bei einer Aufsichtsbehörde",
    f"Sie haben das Recht, sich bei einer Datenschutz-Aufsichtsbehörde zu beschweren, wenn Sie "
    f"der Ansicht sind, dass die Verarbeitung Ihrer Daten gegen die DSGVO verstößt.\n\n"
    f"Zuständige Aufsichtsbehörde: {DATEN['aufsicht_name']}\n"
    f"Adresse: {DATEN['aufsicht_adresse']}\n"
    f"Tel.: {DATEN['aufsicht_tel']}  ·  E-Mail: {DATEN['aufsicht_email']}  ·  {DATEN['aufsicht_web']}")

# ── Fußzeile ──────────────────────────────────────────────────────────────────
doc.add_paragraph().paragraph_format.space_after = Pt(6)
fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_run(fp,
    "Dieses Informationsblatt wurde gemäß Art. 12–13 DSGVO erstellt · "
    "Vorlage: LfD Niedersachsen · Stand: Februar 2026",
    size=8, italic=True, color=RGBColor(0xAA, 0xAA, 0xAA))

doc.save(OUTPUT)
print("Gespeichert:", OUTPUT)
