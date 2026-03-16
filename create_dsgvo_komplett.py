"""
DSGVO-Videoüberwachung – Eiszeit (Marta und Florian Schmidl)
Kombiniertes Dokument:
  Seite 1  – Hinweisschild        (DIN A4, mind. A4 drucken)
  Seite 2  – Informationsblatt    (DIN A3, mind. A3 drucken)
Grundlage: Art. 12–13 DSGVO · Vorlage LfD Niedersachsen
"""
import sys
sys.path.insert(0, r'C:\Users\localUser\AppData\Roaming\Python\Python314\site-packages')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_SECTION

# ── Farben ────────────────────────────────────────────────────────────────────
BEIGE      = RGBColor(0xF0, 0xE4, 0xCB)
BLAU       = RGBColor(0x9B, 0xD4, 0xDB)
GRUEN      = RGBColor(0x6C, 0x94, 0x8C)
SCHRIFT    = RGBColor(0x41, 0x45, 0x45)
WEISS      = RGBColor(0xFF, 0xFF, 0xFF)
DUNKELBLAU = RGBColor(0x1A, 0x5F, 0x7A)
HELLGRAU   = RGBColor(0xF5, 0xF5, 0xF5)

DATEN = {
    "name":      "Marta und Florian Schmidl",
    "adresse":   "Heiligenbruch 27, 27339 Riede",
    "email":     "eiszeit@mailbox.org",
    "zweck":     "Vandalismusprävention, Schutz des Eigentums, Hausrecht",
    "rg":        "Art. 6 Abs. 1 lit. f DSGVO (berechtigtes Interesse)",
    "interesse": "Schutz des Eigentums und der Vertrauenskasse vor Vandalismus und Diebstahl",
    "dauer":     "7 Tage, danach automatische Löschung",
    "dauer_lang":"Die Aufnahmen werden automatisch nach 7 Tagen gelöscht, sofern kein konkreter Vorfall dokumentiert werden muss.",
    "empfaenger":"Die Aufnahmen werden grundsätzlich nicht an Dritte weitergegeben. Im Fall eines dokumentierten Vorfalls können Aufnahmen an Strafverfolgungsbehörden übermittelt werden.",
    "drittland": "Eine Übermittlung personenbezogener Daten an Drittländer oder internationale Organisationen findet nicht statt.",
    "aufsicht":  "Die Landesbeauftragte für den Datenschutz Niedersachsen\nPrinzenstr. 5, 30159 Hannover\nTel.: 0511 120 4500  ·  poststelle@lfd.niedersachsen.de  ·  www.lfd.niedersachsen.de",
}

LOGO   = "EiszeitLogo.png"
KAMERA = "kamera_piktogramm.png"
OUTPUT = "DSGVO_Eiszeit_Komplett_v2.docx"


# ── Hilfsfunktionen ───────────────────────────────────────────────────────────

def cell_bg(cell, rgb):
    tc = cell._tc
    pr = tc.get_or_add_tcPr()
    s  = OxmlElement("w:shd")
    s.set(qn("w:val"),   "clear")
    s.set(qn("w:color"), "auto")
    s.set(qn("w:fill"),  f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    pr.append(s)

def cell_border(cell, color="9BD4DB", sz=8):
    tc = cell._tc
    pr = tc.get_or_add_tcPr()
    b  = OxmlElement("w:tcBorders")
    for side in ("top","bottom","left","right"):
        e = OxmlElement(f"w:{side}")
        e.set(qn("w:val"),   "single")
        e.set(qn("w:sz"),    str(sz))
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), color)
        b.append(e)
    pr.append(b)

def run(para, text, bold=False, italic=False, size=10, color=None):
    r = para.add_run(text)
    r.bold       = bold
    r.italic     = italic
    r.font.name  = "Calibri"
    r.font.size  = Pt(size)
    if color:
        r.font.color.rgb = color
    return r

def gap(doc, pt=3):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(pt)

def header_row(tbl, text, bg=GRUEN, total_cols=2):
    row  = tbl.add_row()
    cell = row.cells[0]
    if total_cols == 2:
        cell.merge(row.cells[1])
    cell_bg(cell, bg)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent  = Cm(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run(p, text, bold=True, size=10, color=WEISS)

def data_row(tbl, label, value, label_w=None):
    row = tbl.add_row()
    lc, rc = row.cells[0], row.cells[1]
    cell_bg(lc, BEIGE);  cell_border(lc)
    cell_bg(rc, WEISS);  cell_border(rc)
    for c, txt, bold, col in ((lc, label, True, GRUEN), (rc, value, False, SCHRIFT)):
        p = c.paragraphs[0]
        p.paragraph_format.left_indent  = Cm(0.3)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        run(p, txt, bold=bold, size=9.5, color=col)

def rights_box(tbl, art, titel, text):
    # Titelzeile
    row1 = tbl.add_row()
    hc   = row1.cells[0]
    hc.merge(row1.cells[1])
    cell_bg(hc, BLAU); cell_border(hc, color="6C948C")
    ph = hc.paragraphs[0]
    ph.paragraph_format.left_indent  = Cm(0.4)
    ph.paragraph_format.space_before = Pt(4)
    ph.paragraph_format.space_after  = Pt(4)
    run(ph, f"{art}  |  {titel}", bold=True, size=10.5, color=DUNKELBLAU)
    # Textzeile
    row2 = tbl.add_row()
    tc2  = row2.cells[0]
    tc2.merge(row2.cells[1])
    cell_bg(tc2, HELLGRAU); cell_border(tc2, color="9BD4DB")
    pt = tc2.paragraphs[0]
    pt.paragraph_format.left_indent  = Cm(0.4)
    pt.paragraph_format.right_indent = Cm(0.4)
    pt.paragraph_format.space_before = Pt(4)
    pt.paragraph_format.space_after  = Pt(5)
    run(pt, text, size=9.5, color=SCHRIFT)

def add_kopf(doc, cam_width, logo_width, col2_width, title_size, sub_size, note_size):
    """Dreispaltiger Kopf: Logo | Kamera | Titeltext"""
    t = doc.add_table(rows=1, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.columns[0].width = logo_width
    t.columns[1].width = cam_width
    t.columns[2].width = col2_width

    lc, cc, tc2 = t.rows[0].cells
    for c in (lc, cc, tc2):
        cell_bg(c, BEIGE)
        cell_border(c, sz=4)

    # Logo
    lp = lc.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lp.paragraph_format.space_before = Pt(8)
    try:
        lp.add_run().add_picture(LOGO, width=logo_width - Cm(0.4))
    except Exception:
        run(lp, "EISZEIT", bold=True, size=14, color=DUNKELBLAU)

    # Kamera
    cp = cc.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(6)
    cp.paragraph_format.space_after  = Pt(6)
    try:
        cp.add_run().add_picture(KAMERA, width=cam_width - Cm(0.6))
    except Exception:
        run(cp, "CCTV", bold=True, size=20, color=DUNKELBLAU)

    # Titel
    p1 = tc2.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p1.paragraph_format.left_indent  = Cm(0.5)
    p1.paragraph_format.space_before = Pt(8)
    run(p1, "Achtung!", bold=True, size=title_size, color=DUNKELBLAU)

    p2 = tc2.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2.paragraph_format.left_indent  = Cm(0.5)
    p2.paragraph_format.space_before = Pt(2)
    run(p2, "Videoüberwachung!", bold=True, size=sub_size, color=SCHRIFT)

    p3 = tc2.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p3.paragraph_format.left_indent  = Cm(0.5)
    p3.paragraph_format.space_before = Pt(5)
    run(p3, "Dieser Bereich wird durch eine Videokamera überwacht.",
        size=note_size, italic=True, color=SCHRIFT)


# ═══════════════════════════════════════════════════════════════════════════════
#  DOKUMENT AUFBAUEN
# ═══════════════════════════════════════════════════════════════════════════════
doc = Document()

# ── SEITE 1: Hinweisschild (DIN A4) ──────────────────────────────────────────
s1 = doc.sections[0]
s1.page_width    = Cm(21.0)
s1.page_height   = Cm(29.7)
s1.top_margin    = Cm(1.5)
s1.bottom_margin = Cm(1.5)
s1.left_margin   = Cm(1.8)
s1.right_margin  = Cm(1.8)

usable = Cm(21.0 - 1.8 - 1.8)   # ~17.4 cm

add_kopf(doc,
    cam_width  = Cm(4.0),
    logo_width = Cm(3.2),
    col2_width = Cm(10.2),
    title_size = 26,
    sub_size   = 18,
    note_size  = 9)

gap(doc, 4)

# Pflichtangaben-Tabelle A4
tbl1 = doc.add_table(rows=0, cols=2)
tbl1.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl1.columns[0].width = Cm(5.0)
tbl1.columns[1].width = Cm(12.4)

header_row(tbl1, "1.  Verantwortlicher (Art. 13 Abs. 1 lit. a DSGVO)")
data_row(tbl1,   "Name:",     DATEN["name"])
data_row(tbl1,   "Adresse:",  DATEN["adresse"])
data_row(tbl1,   "E-Mail:",   DATEN["email"])

header_row(tbl1, "2.  Datenschutzbeauftragter (Art. 13 Abs. 1 lit. b DSGVO)")
data_row(tbl1,   "Hinweis:",  "Kein betrieblicher Datenschutzbeauftragter bestellt.")

header_row(tbl1, "3.  Zweck und Rechtsgrundlage (Art. 13 Abs. 1 lit. c–d DSGVO)")
data_row(tbl1,   "Zweck:",          DATEN["zweck"])
data_row(tbl1,   "Rechtsgrundlage:",DATEN["rg"])
data_row(tbl1,   "Berechtigtes Interesse:", DATEN["interesse"])

header_row(tbl1, "4.  Speicherdauer (Art. 13 Abs. 2 lit. a DSGVO)")
data_row(tbl1,   "Dauer:",   DATEN["dauer"])

header_row(tbl1, "5.  Weitere Informationen", bg=DUNKELBLAU)
data_row(tbl1,   "Informationsblatt:",
    "Das vollständige Informationsblatt mit Ihren Betroffenenrechten "
    "(Auskunft, Widerspruch, Löschung, Beschwerde) finden Sie als "
    "Aushang am Standort (Seite 2 dieses Dokuments).")
data_row(tbl1,   "Aufsichtsbehörde:", DATEN["aufsicht"])

gap(doc, 6)
fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(fp, "Erstellt gemäß Art. 12–13 DSGVO  ·  Vorlage: LfD Niedersachsen  ·  Stand: Februar 2026",
    size=7.5, italic=True, color=RGBColor(0xAA,0xAA,0xAA))


# ── SEITENUMBRUCH + neue Section (ebenfalls A4) ───────────────────────────────
new_sec = doc.add_section(WD_SECTION.NEW_PAGE)
new_sec.page_width    = Cm(21.0)
new_sec.page_height   = Cm(29.7)
new_sec.top_margin    = Cm(1.5)
new_sec.bottom_margin = Cm(1.5)
new_sec.left_margin   = Cm(1.8)
new_sec.right_margin  = Cm(1.8)


# ── SEITE 2: Informationsblatt (DIN A4) ──────────────────────────────────────
add_kopf(doc,
    cam_width  = Cm(4.0),
    logo_width = Cm(3.2),
    col2_width = Cm(10.2),
    title_size = 22,
    sub_size   = 15,
    note_size  = 9)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_before = Pt(4)
run(p_sub,
    "Vollständiges Informationsblatt gemäß Art. 13 DSGVO  ·  Standort: Heiligenbruch 27, 27339 Riede",
    size=10, italic=True, color=SCHRIFT)

gap(doc, 5)

# Inhaltstabelle A4
tbl2 = doc.add_table(rows=0, cols=2)
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl2.columns[0].width = Cm(5.0)
tbl2.columns[1].width = Cm(12.4)

header_row(tbl2, "1.  Verantwortlicher (Art. 13 Abs. 1 lit. a DSGVO)", total_cols=2)
data_row(tbl2, "Name:",     DATEN["name"])
data_row(tbl2, "Adresse:",  DATEN["adresse"])
data_row(tbl2, "E-Mail:",   DATEN["email"])

header_row(tbl2, "2.  Datenschutzbeauftragter (Art. 13 Abs. 1 lit. b DSGVO)", total_cols=2)
data_row(tbl2, "Hinweis:",
    "Es wurde kein betrieblicher Datenschutzbeauftragter bestellt, "
    "da die Voraussetzungen gemäß Art. 37 DSGVO nicht vorliegen.")

header_row(tbl2, "3.  Zweck und Rechtsgrundlage (Art. 13 Abs. 1 lit. c–d DSGVO)", total_cols=2)
data_row(tbl2, "Zweck:",            DATEN["zweck"])
data_row(tbl2, "Rechtsgrundlage:",  DATEN["rg"])
data_row(tbl2, "Berechtigtes\nInteresse:", DATEN["interesse"])

header_row(tbl2, "4.  Speicherdauer (Art. 13 Abs. 2 lit. a DSGVO)", total_cols=2)
data_row(tbl2, "Dauer:", DATEN["dauer_lang"])

header_row(tbl2, "5.  Empfänger der Daten (Art. 13 Abs. 1 lit. e DSGVO)", total_cols=2)
data_row(tbl2, "Empfänger:",   DATEN["empfaenger"])
data_row(tbl2, "Drittländer:", DATEN["drittland"])

header_row(tbl2, "6.  Ihre Rechte als betroffene Person", bg=DUNKELBLAU, total_cols=2)
rights_box(tbl2, "Art. 15 DSGVO", "Recht auf Auskunft",
    "Sie haben das Recht, von uns eine Bestätigung darüber zu verlangen, ob Sie betreffende "
    "personenbezogene Daten verarbeitet werden. Ist dies der Fall, haben Sie ein Recht auf "
    "Auskunft über diese Daten sowie auf die in Art. 15 DSGVO aufgeführten Informationen.")
rights_box(tbl2, "Art. 16 DSGVO", "Recht auf Berichtigung",
    "Sie haben das Recht, von uns unverzüglich die Berichtigung Sie betreffender unrichtiger "
    "personenbezogener Daten sowie ggf. die Vervollständigung unvollständiger Daten zu verlangen.")
rights_box(tbl2, "Art. 17 DSGVO", "Recht auf Löschung",
    "Sie haben das Recht, von uns zu verlangen, dass Sie betreffende personenbezogene Daten "
    "unverzüglich gelöscht werden, sofern einer der in Art. 17 DSGVO genannten Gründe zutrifft – "
    "z. B. wenn die Daten für die verfolgten Zwecke nicht mehr benötigt werden.")
rights_box(tbl2, "Art. 18 DSGVO", "Recht auf Einschränkung der Verarbeitung",
    "Sie haben das Recht, von uns die Einschränkung der Verarbeitung zu verlangen, wenn eine der "
    "in Art. 18 DSGVO aufgeführten Voraussetzungen gegeben ist – z. B. wenn Sie Widerspruch "
    "gegen die Verarbeitung eingelegt haben, für die Dauer der Prüfung durch uns.")
rights_box(tbl2, "Art. 21 DSGVO", "Recht auf Widerspruch",
    "Sie haben das Recht, aus Gründen, die sich aus Ihrer besonderen Situation ergeben, jederzeit "
    "Widerspruch gegen die Verarbeitung Sie betreffender personenbezogener Daten einzulegen. "
    "Wir verarbeiten die Daten dann nicht mehr, es sei denn, wir können zwingende schutzwürdige "
    "Gründe nachweisen, die Ihre Interessen überwiegen.")
rights_box(tbl2, "Art. 77 DSGVO", "Recht auf Beschwerde bei einer Aufsichtsbehörde",
    "Sie haben das Recht, sich bei einer Datenschutz-Aufsichtsbehörde zu beschweren, wenn Sie "
    "der Ansicht sind, dass die Verarbeitung Ihrer Daten gegen die DSGVO verstößt.\n\n"
    "Zuständige Aufsichtsbehörde:\n" + DATEN["aufsicht"])

gap(doc, 6)
fp2 = doc.add_paragraph()
fp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(fp2,
    "Erstellt gemäß Art. 12–13 DSGVO  ·  Vorlage: LfD Niedersachsen  ·  Stand: Februar 2026",
    size=8, italic=True, color=RGBColor(0xAA,0xAA,0xAA))

doc.save(OUTPUT)
print("Gespeichert:", OUTPUT)
